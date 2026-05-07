#!/usr/bin/env python3
"""Generate a final Chinese DOCX report for Happy Monkey.

This version integrates:
1) distance cleanup summary (legacy RDC rows removed)
2) data quality acceptance checks
3) dual-city model rerun based on final 3 tables only (no builtin merge)

Typography/layout rules:
- Body font: 宋体, 五号 (10.5pt)
- Line spacing: 1.5
- Paragraph spacing: 0 before/after
- First-line indent: 2 chars (~21pt)
- Figure captions below images, table captions above tables
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import shutil
import sys
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Sequence, Tuple

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app import parse_csv_file
from optimizer import city_params, run_full_model


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _safe_float(value, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return default


def load_rows(path: Path) -> List[Dict[str, object]]:
    return parse_csv_file(path)


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def cleanup_distances(distances_path: Path, rdcs_path: Path) -> Dict[str, object]:
    """Backup + filter distances rows whose rdc_id is not in rdcs.csv."""
    rdcs_rows = load_rows(rdcs_path)
    valid_rdc = {str(r.get("rdc_id") or "").strip() for r in rdcs_rows if str(r.get("rdc_id") or "").strip()}

    with distances_path.open(newline="", encoding="utf-8-sig") as f:
        distance_rows = list(csv.DictReader(f))

    fields = list(distance_rows[0].keys()) if distance_rows else ["rdc_id", "store_id", "distance_km"]

    removed: List[Dict[str, str]] = []
    kept: List[Dict[str, str]] = []
    for row in distance_rows:
        rid = str(row.get("rdc_id") or "").strip()
        if rid in valid_rdc:
            kept.append(row)
        else:
            removed.append(row)

    backup_path = distances_path.with_name(f"{distances_path.stem}_backup_{_timestamp()}{distances_path.suffix}")
    shutil.copy2(distances_path, backup_path)

    with distances_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(kept)

    removed_by_rdc = dict(Counter(str(r.get("rdc_id") or "").strip() for r in removed))

    return {
        "cleaned": True,
        "backup_path": str(backup_path),
        "original_rows": len(distance_rows),
        "kept_rows": len(kept),
        "removed_rows": len(removed),
        "removed_by_rdc": removed_by_rdc,
        "cleaned_at": datetime.now().isoformat(timespec="seconds"),
    }


def derive_cleanup_summary_from_backup(backup_path: Path, cleaned_path: Path, rdcs_path: Path) -> Dict[str, object]:
    rdcs_rows = load_rows(rdcs_path)
    valid_rdc = {str(r.get("rdc_id") or "").strip() for r in rdcs_rows if str(r.get("rdc_id") or "").strip()}

    with backup_path.open(newline="", encoding="utf-8-sig") as f:
        original_rows = list(csv.DictReader(f))
    with cleaned_path.open(newline="", encoding="utf-8-sig") as f:
        cleaned_rows = list(csv.DictReader(f))

    removed = [r for r in original_rows if str(r.get("rdc_id") or "").strip() not in valid_rdc]

    return {
        "cleaned": False,
        "backup_path": str(backup_path),
        "original_rows": len(original_rows),
        "kept_rows": len(cleaned_rows),
        "removed_rows": len(removed),
        "removed_by_rdc": dict(Counter(str(r.get("rdc_id") or "").strip() for r in removed)),
        "cleaned_at": datetime.now().isoformat(timespec="seconds"),
    }


def compute_quality_checks(
    stores_rows: Sequence[Dict[str, object]],
    rdcs_rows: Sequence[Dict[str, object]],
    distance_rows: Sequence[Dict[str, object]],
) -> Dict[str, object]:
    store_ids = [str(r.get("store_id") or "").strip() for r in stores_rows]
    rdc_ids = [str(r.get("rdc_id") or "").strip() for r in rdcs_rows]
    store_set = {s for s in store_ids if s}
    rdc_set = {r for r in rdc_ids if r}

    unknown_rdc_rows = [r for r in distance_rows if str(r.get("rdc_id") or "").strip() not in rdc_set]
    unknown_store_rows = [r for r in distance_rows if str(r.get("store_id") or "").strip() not in store_set]

    store_city = {str(r.get("store_id") or "").strip(): str(r.get("city") or "").strip().lower() for r in stores_rows}
    rdc_city = {str(r.get("rdc_id") or "").strip(): str(r.get("city") or "").strip().lower() for r in rdcs_rows}

    covered = defaultdict(int)
    by_store = defaultdict(set)
    distance_values = []

    for r in distance_rows:
        rid = str(r.get("rdc_id") or "").strip()
        sid = str(r.get("store_id") or "").strip()
        if rid in rdc_city and sid in store_city:
            covered[(rdc_city[rid], store_city[sid])] += 1
            by_store[sid].add(rid)
        try:
            distance_values.append(float(r.get("distance_km") or 0.0))
        except Exception:
            pass

    same_city_coverage = {}
    for city in ["beijing", "hangzhou"]:
        n_store = sum(1 for r in stores_rows if str(r.get("city") or "").strip().lower() == city)
        n_rdc = sum(1 for r in rdcs_rows if str(r.get("city") or "").strip().lower() == city)
        expected = n_store * n_rdc
        actual = covered.get((city, city), 0)
        pct = (actual / expected * 100.0) if expected else 0.0
        same_city_coverage[city] = {
            "expected": expected,
            "actual": actual,
            "coverage_pct": round(pct, 2),
        }

    existing_stores = [r for r in stores_rows if str(r.get("is_existing") or "").strip() == "1"]
    missing_existing_pairs = []
    for r in existing_stores:
        sid = str(r.get("store_id") or "").strip()
        city = str(r.get("city") or "").strip().lower()
        city_rdcs = [rid for rid, c in rdc_city.items() if c == city]
        miss = [rid for rid in city_rdcs if rid not in by_store.get(sid, set())]
        if miss:
            missing_existing_pairs.append({"store_id": sid, "missing_rdcs": miss})

    distance_stats = {
        "min": round(min(distance_values), 4) if distance_values else None,
        "max": round(max(distance_values), 4) if distance_values else None,
        "avg": round(sum(distance_values) / len(distance_values), 4) if distance_values else None,
        "negative_count": sum(1 for v in distance_values if v < 0),
        "zero_count": sum(1 for v in distance_values if v == 0),
    }

    return {
        "distance_unknown_rdc_rows": len(unknown_rdc_rows),
        "distance_unknown_store_rows": len(unknown_store_rows),
        "same_city_coverage": same_city_coverage,
        "existing_store_missing_pairs_count": len(missing_existing_pairs),
        "existing_store_missing_pairs": missing_existing_pairs,
        "duplicate_store_ids": [k for k, v in Counter(store_ids).items() if k and v > 1],
        "duplicate_rdc_ids": [k for k, v in Counter(rdc_ids).items() if k and v > 1],
        "distance_stats": distance_stats,
        "stores_count": len(stores_rows),
        "rdcs_count": len(rdcs_rows),
        "distances_count": len(distance_rows),
    }


def configure_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    p = normal.paragraph_format
    p.line_spacing = 1.5
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    p.first_line_indent = Pt(21)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "宋体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing = 1.5

    h2 = doc.styles["Heading 2"]
    h2.font.name = "宋体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.line_spacing = 1.5

    if "CaptionCN" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("CaptionCN", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["CaptionCN"]
    cap.font.name = "宋体"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(10.5)
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.line_spacing = 1.5


class CaptionCounter:
    def __init__(self) -> None:
        self.figure_idx = 0
        self.table_idx = 0

    def figure(self, title: str) -> str:
        self.figure_idx += 1
        return f"图 {self.figure_idx} {title}"

    def table(self, title: str) -> str:
        self.table_idx += 1
        return f"表 {self.table_idx} {title}"


def _set_run_songti(run, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]


def add_figure(doc: Document, image_path: Path, caption: str, max_width_inch: float = 6.3) -> None:
    fig_para = doc.add_paragraph()
    fig_para.paragraph_format.first_line_indent = Pt(0)
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(str(image_path), width=Inches(max_width_inch))

    cap = doc.add_paragraph(caption)
    cap.style = doc.styles["CaptionCN"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _format_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(str(text))
    _set_run_songti(run, size_pt=10.5, bold=bold)


def add_table(doc: Document, caption: str, headers: Sequence[str], rows: Sequence[Sequence[object]]) -> None:
    cap = doc.add_paragraph(caption)
    cap.style = doc.styles["CaptionCN"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    for j, header in enumerate(headers):
        _format_cell(table.rows[0].cells[j], str(header), bold=True)

    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            _format_cell(table.rows[i].cells[j], str(value), bold=False)


def build_datasets(stores_path: Path, rdcs_path: Path, distances_path: Path) -> Dict[str, object]:
    stores_rows = load_rows(stores_path)
    rdcs_rows = load_rows(rdcs_path)
    distance_rows = load_rows(distances_path)

    stores_df = pd.DataFrame(stores_rows)
    rdcs_df = pd.DataFrame(rdcs_rows)
    dist_df = pd.DataFrame(distance_rows)

    for df, col in [(stores_df, "city"), (rdcs_df, "city")]:
        if col in df.columns:
            df[col] = df[col].astype(str).str.lower().str.strip()

    result_beijing = run_full_model(
        store_rows=stores_rows,
        rdc_rows=rdcs_rows,
        distance_rows=distance_rows,
        max_new_stores=20,
        p_values=[1, 2, 3, 4, 5],
        focus_city="beijing",
        threshold_overrides=None,
    )
    result_hangzhou = run_full_model(
        store_rows=stores_rows,
        rdc_rows=rdcs_rows,
        distance_rows=distance_rows,
        max_new_stores=20,
        p_values=[1, 2, 3, 4, 5],
        focus_city="hangzhou",
        threshold_overrides=None,
    )

    checks = compute_quality_checks(stores_rows, rdcs_rows, distance_rows)

    return {
        "stores_df": stores_df,
        "rdcs_df": rdcs_df,
        "dist_df": dist_df,
        "result_beijing": result_beijing,
        "result_hangzhou": result_hangzhou,
        "store_rows": stores_rows,
        "rdc_rows": rdcs_rows,
        "distance_rows": distance_rows,
        "quality_checks": checks,
    }


def _setup_plot_font() -> None:
    plt.rcParams["font.sans-serif"] = ["Songti SC", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def generate_figures(payload: Dict[str, object], figure_dir: Path) -> Dict[str, Path]:
    ensure_dir(figure_dir)
    _setup_plot_font()

    stores_df: pd.DataFrame = payload["stores_df"]
    rdcs_df: pd.DataFrame = payload["rdcs_df"]
    res_bj = payload["result_beijing"]
    res_hz = payload["result_hangzhou"]
    checks = payload["quality_checks"]

    output: Dict[str, Path] = {}

    fig1 = figure_dir / "fig1_store_city_type_count.png"
    ct = (
        stores_df.assign(type=stores_df["is_existing"].apply(lambda x: "existing" if int(_safe_float(x)) == 1 else "new"))
        .groupby(["city", "type"])  # type: ignore[arg-type]
        .size()
        .unstack(fill_value=0)
    )
    ct = ct.reindex(index=["beijing", "hangzhou"], fill_value=0)
    ax = ct.plot(kind="bar", figsize=(7, 4), color=["#4c78a8", "#f58518"])
    ax.set_title("Store composition by city")
    ax.set_xlabel("City")
    ax.set_ylabel("Count")
    ax.legend(title="Type")
    plt.tight_layout()
    plt.savefig(fig1, dpi=180)
    plt.close()
    output["store_city_type"] = fig1

    fig2 = figure_dir / "fig2_stage1_pass_rate.png"
    summary_rows = []
    for city, res in [("beijing", res_bj), ("hangzhou", res_hz)]:
        s = res["summary"]
        new_count = max(int(s.get("new_store_candidates", 0)), 1)
        pass_count = int(s.get("stage1_passed", 0))
        summary_rows.append((city, pass_count / new_count * 100.0, pass_count, new_count))
    pass_df = pd.DataFrame(summary_rows, columns=["city", "pass_rate", "pass_count", "new_count"])
    ax = pass_df.plot(x="city", y="pass_rate", kind="bar", legend=False, figsize=(6.5, 4), color="#54a24b")
    ax.set_title("Stage-1 pass rate by city")
    ax.set_xlabel("City")
    ax.set_ylabel("Pass rate (%)")
    for i, row in pass_df.iterrows():
        ax.text(i, row["pass_rate"] + 1.0, f"{row['pass_count']}/{row['new_count']}", ha="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(fig2, dpi=180)
    plt.close()
    output["stage1_pass_rate"] = fig2

    fig3 = figure_dir / "fig3_stage3_cost_breakdown.png"
    cost_rows = []
    for city, res in [("beijing", res_bj), ("hangzhou", res_hz)]:
        best = res["stage3"].get("best_scenario") or {}
        cost_rows.append(
            {
                "city": city,
                "rdc_cost_10k": _safe_float(best.get("rdc_cost_10k", 0.0)),
                "delivery_cost_10k": _safe_float(best.get("delivery_cost_10k", 0.0)),
            }
        )
    cost_df = pd.DataFrame(cost_rows).set_index("city")
    ax = cost_df.plot(kind="bar", stacked=True, figsize=(7, 4), color=["#4c78a8", "#e45756"])
    ax.set_title("Stage-3 cost components by city")
    ax.set_xlabel("City")
    ax.set_ylabel("Cost (10k CNY, PV)")
    ax.legend(title="Component")
    plt.tight_layout()
    plt.savefig(fig3, dpi=180)
    plt.close()
    output["stage3_cost"] = fig3

    fig4 = figure_dir / "fig4_best_p_objective.png"
    comp_rows = []
    for city, res in [("beijing", res_bj), ("hangzhou", res_hz)]:
        s = res["summary"]
        comp_rows.append({"city": city, "best_p": _safe_float(s.get("best_p", 0.0)), "objective_10k": _safe_float(s.get("best_joint_objective_10k", 0.0))})
    comp_df = pd.DataFrame(comp_rows)
    fig, ax1 = plt.subplots(figsize=(7, 4))
    ax1.bar(comp_df["city"], comp_df["objective_10k"], color="#72b7b2", label="joint objective")
    ax1.set_ylabel("Joint objective (10k)")
    ax1.set_xlabel("City")
    ax1.set_title("Best P and objective by city")
    ax2 = ax1.twinx()
    ax2.plot(comp_df["city"], comp_df["best_p"], color="#f58518", marker="o", linewidth=2, label="best P")
    ax2.set_ylabel("Best P")
    fig.tight_layout()
    plt.savefig(fig4, dpi=180)
    plt.close(fig)
    output["best_p_objective"] = fig4

    fig5 = figure_dir / "fig5_spatial_distribution.png"
    fig, ax = plt.subplots(figsize=(7, 5))
    for city, color in [("beijing", "#4c78a8"), ("hangzhou", "#f58518")]:
        sub_s = stores_df[stores_df["city"] == city]
        sub_r = rdcs_df[rdcs_df["city"] == city]
        if not sub_s.empty:
            ax.scatter(sub_s["lon"], sub_s["lat"], s=24, alpha=0.7, c=color, marker="o", label=f"{city}-stores")
        if not sub_r.empty:
            ax.scatter(sub_r["lon"], sub_r["lat"], s=52, alpha=0.9, c=color, marker="^", label=f"{city}-rdcs")
    ax.set_title("Spatial distribution of stores and RDC candidates")
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    ax.legend(fontsize=8, ncol=2)
    plt.tight_layout()
    plt.savefig(fig5, dpi=180)
    plt.close(fig)
    output["spatial_distribution"] = fig5

    fig6 = figure_dir / "fig6_same_city_coverage.png"
    cov_rows = []
    for city in ["beijing", "hangzhou"]:
        row = checks["same_city_coverage"].get(city, {})
        cov_rows.append({"city": city, "coverage_pct": _safe_float(row.get("coverage_pct", 0.0))})
    cov_df = pd.DataFrame(cov_rows)
    ax = cov_df.plot(x="city", y="coverage_pct", kind="bar", legend=False, figsize=(6.5, 4), color="#9c755f")
    ax.set_title("Same-city distance coverage")
    ax.set_xlabel("City")
    ax.set_ylabel("Coverage (%)")
    ax.set_ylim(0, 105)
    for i, row in cov_df.iterrows():
        ax.text(i, row["coverage_pct"] + 1.0, f"{row['coverage_pct']:.2f}%", ha="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(fig6, dpi=180)
    plt.close()
    output["same_city_coverage"] = fig6

    return output


def _build_table_field_desc() -> Dict[str, List[List[str]]]:
    stores_rows = [
        ["store_id", "门店唯一ID", "string"],
        ["name", "门店名称", "string"],
        ["city", "城市（beijing/hangzhou）", "string"],
        ["lat/lon", "坐标（WGS84）", "float"],
        ["rf_sales_10k", "参考年销售额（万元）", "float"],
        ["area_sqm", "门店面积（㎡）", "float"],
        ["initial_investment_10k", "初始投资（万元）", "float"],
        ["annual_fixed_cost_10k", "年固定成本（万元）", "float"],
        ["is_existing", "是否既有门店（1是/0否）", "int"],
    ]

    rdcs_rows = [
        ["rdc_id", "RDC唯一ID", "string"],
        ["name", "RDC名称", "string"],
        ["city", "城市", "string"],
        ["lat/lon", "坐标（WGS84）", "float"],
        ["area_sqm", "仓库面积（㎡）", "float"],
        ["initial_investment_10k", "初始投资（万元）", "float"],
        ["annual_*_10k", "年租金/物业/运营/人工/水电（万元）", "float"],
        ["residual_rate", "残值率", "float"],
        ["rent_per_sqm_day", "租金（元/㎡/天）", "float"],
        ["lease_years", "剩余租期（年）", "float"],
        ["blue_access", "蓝牌通行", "int/bool"],
        ["core_travel_min", "到核心区通勤时间（分钟）", "float"],
        ["clear_height_m", "净高（米）", "float"],
        ["loading_dock", "是否有装卸月台", "int/bool"],
        ["can_three_temp", "是否支持三温仓", "int/bool"],
        ["fire_pass", "消防是否通过", "int/bool"],
        ["disturbance_risk", "扰民风险（1有风险）", "int/bool"],
    ]

    dist_rows = [
        ["rdc_id", "RDC编号", "string"],
        ["store_id", "门店编号", "string"],
        ["distance_km", "距离（公里）", "float"],
    ]

    return {"stores": stores_rows, "rdcs": rdcs_rows, "distances": dist_rows}


def _city_param_rows() -> List[List[object]]:
    rows = []
    for city in ["beijing", "hangzhou"]:
        cp = city_params(city)
        rows.append(
            [
                city,
                cp["discount_rate"],
                cp["gross_margin"],
                cp["variable_cost_rate"],
                cp["tax_rate"],
                cp["npv_threshold_10k"],
                cp["dpp_threshold_years"],
                cp["revenue_per_sqm_threshold_10k"],
                cp["rent_cap_per_sqm_day"],
            ]
        )
    return rows


def _result_summary_rows(res_bj: Dict[str, object], res_hz: Dict[str, object]) -> List[List[object]]:
    rows = []
    for city, res in [("beijing", res_bj), ("hangzhou", res_hz)]:
        s = res["summary"]
        best = res["stage3"].get("best_scenario") or {}
        rows.append(
            [
                city,
                s.get("new_store_candidates", 0),
                s.get("stage1_passed", 0),
                s.get("stage2_selected", 0),
                s.get("best_p", "-"),
                best.get("rdc_cost_10k", "-"),
                best.get("delivery_cost_10k", "-"),
                best.get("total_cost_10k", "-"),
                best.get("roi", "-"),
            ]
        )
    return rows


def _open_rdc_rows(city: str, res: Dict[str, object]) -> List[List[object]]:
    best = res["stage3"].get("best_scenario") or {}
    out = []
    for r in best.get("open_rdcs", []) or []:
        out.append([city, r.get("rdc_id"), r.get("name"), r.get("lifecycle_cost_10k")])
    return out


def build_docx(payload: Dict[str, object], figures: Dict[str, Path], cleanup_summary: Dict[str, object], out_doc: Path) -> None:
    res_bj = payload["result_beijing"]
    res_hz = payload["result_hangzhou"]
    stores_df: pd.DataFrame = payload["stores_df"]
    rdcs_df: pd.DataFrame = payload["rdcs_df"]
    dist_df: pd.DataFrame = payload["dist_df"]
    checks = payload["quality_checks"]

    doc = Document()
    configure_document_styles(doc)
    counter = CaptionCounter()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("快乐猴社区超市联合选址优化项目报告（最终三表版）")
    _set_run_songti(run, size_pt=22, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Pt(0)
    run2 = p2.add_run("（数据清理 + 双城结果 + 正式版排版）")
    _set_run_songti(run2, size_pt=14, bold=False)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.first_line_indent = Pt(0)
    run3 = p3.add_run(datetime.now().strftime("生成日期：%Y-%m-%d"))
    _set_run_songti(run3, size_pt=12, bold=False)

    doc.add_page_break()

    doc.add_heading("一、数据的说明和描述", level=1)
    add_body_paragraph(
        doc,
        "本节对最终三表（stores.csv、rdcs.csv、distances.csv）逐字段说明，并给出数据清理与质量验收结果。"
        "本报告计算口径固定为仅使用最终三表，不并入内置既有门店。"
    )

    add_table(
        doc,
        counter.table("输入数据文件规模统计"),
        ["文件名", "记录数", "字段数", "说明"],
        [
            ["stores.csv", len(stores_df), len(stores_df.columns), "门店候选点与既有门店"],
            ["rdcs.csv", len(rdcs_df), len(rdcs_df.columns), "RDC候选点"],
            ["distances.csv", len(dist_df), len(dist_df.columns), "RDC-门店距离矩阵"],
        ],
    )

    add_table(
        doc,
        counter.table("距离表清理摘要（旧RDC记录剔除）"),
        ["指标", "数值"],
        [
            ["清理前行数", cleanup_summary.get("original_rows")],
            ["清理后行数", cleanup_summary.get("kept_rows")],
            ["删除行数", cleanup_summary.get("removed_rows")],
            ["删除明细", json.dumps(cleanup_summary.get("removed_by_rdc", {}), ensure_ascii=False)],
            ["备份文件", cleanup_summary.get("backup_path", "-")],
            ["清理时间", cleanup_summary.get("cleaned_at", "-")],
        ],
    )

    add_table(
        doc,
        counter.table("数据验收检查结果"),
        ["检查项", "结果"],
        [
            ["distance_unknown_rdc_rows", checks.get("distance_unknown_rdc_rows")],
            ["distance_unknown_store_rows", checks.get("distance_unknown_store_rows")],
            ["北京同城覆盖(实际/应有)", f"{checks['same_city_coverage']['beijing']['actual']}/{checks['same_city_coverage']['beijing']['expected']}"],
            ["杭州同城覆盖(实际/应有)", f"{checks['same_city_coverage']['hangzhou']['actual']}/{checks['same_city_coverage']['hangzhou']['expected']}"],
            ["既有门店缺失对数", checks.get("existing_store_missing_pairs_count")],
        ],
    )

    field_desc = _build_table_field_desc()
    add_table(doc, counter.table("stores.csv 字段说明"), ["字段", "含义", "类型"], field_desc["stores"])
    add_table(doc, counter.table("rdcs.csv 字段说明"), ["字段", "含义", "类型"], field_desc["rdcs"])
    add_table(doc, counter.table("distances.csv 字段说明"), ["字段", "含义", "类型"], field_desc["distances"])

    add_table(
        doc,
        counter.table("模型核心城市参数（代码默认值）"),
        [
            "城市",
            "discount_rate",
            "gross_margin",
            "variable_cost_rate",
            "tax_rate",
            "npv_threshold_10k",
            "dpp_threshold_years",
            "revenue_per_sqm_threshold_10k",
            "rent_cap_per_sqm_day",
        ],
        _city_param_rows(),
    )

    doc.add_heading("二、数据的分析与解读", level=1)
    add_body_paragraph(
        doc,
        "双城样本结构与空间分布反映了不同城市圈层对选址策略的影响。"
        "本节通过结构图与覆盖率图展示数据基础，并明确外部来源引用。"
    )

    add_figure(doc, figures["store_city_type"], counter.figure("双城门店结构分布（既有/新增）"))
    add_figure(doc, figures["spatial_distribution"], counter.figure("门店与RDC候选点空间分布（经纬度散点）"))
    add_figure(doc, figures["stage1_pass_rate"], counter.figure("Stage-1 财务筛选通过率对比"))
    add_figure(doc, figures["same_city_coverage"], counter.figure("同城距离矩阵覆盖率"))

    add_body_paragraph(
        doc,
        "数据来源说明：门店、RDC与距离矩阵来自最终三表；道路距离API与政策说明用于业务解释。"
        "参考文献采用GB/T 7714编号格式。"
    )

    doc.add_heading("三、模型的建立和求解", level=1)
    add_body_paragraph(
        doc,
        "本项目采用三阶段联合建模框架。第一阶段解决单店财务可行性筛选，第二阶段解决新增门店间的空间竞争与蚕食，"
        "第三阶段解决RDC选址与门店分配问题。为保证说明与程序一致，以下数学模型均按当前代码实现方式展开。"
    )

    doc.add_heading("3.1 符号、集合与参数定义", level=2)
    add_table(
        doc,
        counter.table("核心集合与索引定义"),
        ["符号", "含义"],
        [
            ["I", "门店全集，含新增候选门店与既有门店"],
            ["N", "新增候选门店集合，N subset I"],
            ["E", "既有门店集合，E subset I"],
            ["J", "RDC候选点集合"],
            ["J_c", "城市 c 下的合格 RDC 集合"],
            ["i, k", "门店索引"],
            ["j", "RDC索引"],
            ["t", "年份索引，t = 1,...,T"],
            ["c", "城市索引，c in {beijing, hangzhou}"],
            ["P", "开放 RDC 数量"],
        ],
    )

    add_table(
        doc,
        counter.table("核心参数与决策变量定义"),
        ["符号", "含义"],
        [
            ["S_i", "门店 i 的基准年销售额（万元）"],
            ["A_i", "门店 i 的面积（平方米）"],
            ["I_i", "门店 i 的初始投资（万元）"],
            ["F_i", "门店 i 的年固定成本（万元）"],
            ["g_t", "第 t 年销售成长系数"],
            ["r_c", "城市 c 的折现率"],
            ["m_c", "城市 c 的毛利率"],
            ["v_c", "城市 c 的变动成本率"],
            ["tau_c", "城市 c 的有效税率"],
            ["rho_i", "门店 i 的残值率"],
            ["d_ik", "门店 i 与门店 k 的距离（km）"],
            ["D_ij", "RDC j 至门店 i 的距离（km）"],
            ["x_j", "若 RDC j 开放则 x_j=1，否则为0"],
            ["z_ij", "若门店 i 分配给 RDC j 则 z_ij=1，否则为0"],
            ["y_i", "若新增门店 i 被选中则 y_i=1，否则为0"],
            ["alpha_i", "门店 i 在蚕食作用下的销售保持系数"],
        ],
    )

    doc.add_heading("3.2 第一阶段：新增门店财务可行性筛选模型", level=2)
    add_body_paragraph(
        doc,
        "第一阶段仅对新增候选门店 i in N 进行逐点财务评估。既有门店 i in E 不参与该阶段筛选，而是在第三阶段直接进入配送网络。"
    )
    add_body_paragraph(doc, "第 t 年营业收入定义为：Revenue_i,t = S_i * g_t。")
    add_body_paragraph(doc, "第 t 年毛利定义为：Gross_i,t = Revenue_i,t * m_c。")
    add_body_paragraph(doc, "第 t 年变动成本定义为：VarCost_i,t = Revenue_i,t * v_c。")
    add_body_paragraph(doc, "第 t 年营业利润定义为：OpProfit_i,t = Gross_i,t - VarCost_i,t - F_i。")
    add_body_paragraph(doc, "第 t 年所得税定义为：Tax_i,t = max(0, OpProfit_i,t) * tau_c。")
    add_body_paragraph(
        doc,
        "第 t 年现金流定义为：CF_i,t = OpProfit_i,t - Tax_i,t；若 t = T，则在 CF_i,T 中额外加入残值 I_i * rho_i。"
    )
    add_body_paragraph(doc, "净现值定义为：NPV_i = -I_i + sum_t [ CF_i,t / (1 + r_c)^t ]。")
    add_body_paragraph(
        doc,
        "折现回收期 DPP_i 采用累计折现现金流首次转正的时点，并用线性插值计算年内回收位置。"
    )
    add_body_paragraph(doc, "坪效指标定义为：Eff_i = S_i / A_i。")
    add_body_paragraph(
        doc,
        "据此，第一阶段判定规则为：若门店 i 同时满足 NPV_i >= NPV_th(c)，DPP_i <= DPP_th(c)，Eff_i >= Eff_th(c)，"
        "则记为通过；否则淘汰。"
    )
    add_body_paragraph(
        doc,
        "因此，第一阶段本质上是一个逐点筛选模型，而非门店间耦合优化模型，其输出是可进入第二阶段的候选集合 N_pass。"
    )

    doc.add_heading("3.3 第二阶段：新增门店蚕食效应与组合优化模型", level=2)
    add_body_paragraph(
        doc,
        "第二阶段在 N_pass 上求解新增门店组合优化问题，核心思想是把门店间距离映射为竞争惩罚，并在惩罚后重算门店价值。"
    )
    add_body_paragraph(
        doc,
        "对于任意两个新增门店 i, k，设其球面距离为 d_ik。代码中的分段蚕食函数 h(d_ik) 定义为："
        "当 d_ik < 0.28 时，h = 1.0；当 0.28 <= d_ik < 0.32 时，h 从 1.0 线性下降到 0.35；"
        "当 0.32 <= d_ik < 0.50 时，h 从 0.35 线性下降到 0；当 d_ik >= 0.50 时，h = 0。"
    )
    add_body_paragraph(
        doc,
        "若 h(d_ik) = 1.0，则门店 i 与门店 k 构成硬冲突，不允许同时被选；若 0 < h(d_ik) < 1，则其进入软惩罚集合。"
    )
    add_body_paragraph(
        doc,
        "对任一被选门店 i，其总惩罚记为 Pen_i = min(0.45, sum_k!=i [ h(d_ik) * y_k ])，"
        "对应销售保持系数 alpha_i = 1 - Pen_i。"
    )
    add_body_paragraph(doc, "蚕食后的销售额定义为：S'_i = alpha_i * S_i。")
    add_body_paragraph(
        doc,
        "在程序实现中，蚕食仅作用于经营性现金流，不作用于初始投资与残值，因此调整后的净现值定义为："
        "AdjNPV_i = -I_i + sum_t [ CF'_i,t / (1 + r_c)^t ]，其中 CF'_i,t 由调整后销售额 S'_i 重新计算。"
    )
    add_body_paragraph(
        doc,
        "同理，调整后的折现回收期记为 AdjDPP_i。若某门店在蚕食后不再满足 AdjNPV_i >= NPV_th(c) 或 AdjDPP_i <= DPP_th(c)，"
        "则该组合被直接判为不可行。"
    )
    add_body_paragraph(
        doc,
        "因此，第二阶段组合优化模型可以表述为：max sum_i in N_pass [ AdjNPV_i * y_i ]。"
        "约束包括：y_i + y_k <= 1（对所有硬冲突对 (i,k)）、sum_i y_i <= M、以及蚕食后的财务可行性约束。"
    )
    add_body_paragraph(
        doc,
        "其中 M 为最大新增门店数。本项目对可行组合的搜索采用两级策略：当理论组合数不超过 250000 时采用精确枚举；"
        "超过该阈值时切换为束搜索，以控制组合爆炸。"
    )

    doc.add_heading("3.4 第三阶段：RDC选址与门店分配模型", level=2)
    add_body_paragraph(
        doc,
        "第三阶段在既有门店与第二阶段选中的新增门店共同构成的网络上，求解 RDC 选址与门店分配问题。"
        "求解前先对 RDC 候选点做硬约束筛选，仅保留满足通行、净高、租金、租期、消防、扰民风险等条件的合格集合 J_c。"
    )
    add_body_paragraph(
        doc,
        "对任一 RDC j，其全生命周期折现成本定义为：C_rdc(j) = Init_j + sum_t [ Annual_j / (1 + r_c)^t ] - Residual_j / (1 + r_c)^T。"
    )
    add_body_paragraph(
        doc,
        "其中 Annual_j 为年租金、物业费、运营费、人工费和水电费之和，Residual_j = Init_j * residual_rate_j。"
    )
    add_body_paragraph(
        doc,
        "对任一门店 i 分配至 RDC j，其单次配送成本定义为：TripCost_ij = f_c + u_c * D_ij + AreaHandle_i。"
        "其中 f_c 为固定配送成本，u_c 为单位里程变动成本，AreaHandle_i 用门店面积弹性近似处理作业成本。"
    )
    add_body_paragraph(
        doc,
        "进一步地，代码按销售规模和面积构造需求权重 DemandWeight_i，并将年配送频次与成长系数一并折现，得到："
        "C_del(i,j) = sum_t [ Freq_c * DemandWeight_i * TripCost_ij * g_t / (1 + r_c)^t ]。"
    )
    add_body_paragraph(
        doc,
        "于是，在给定 P 的情况下，第三阶段可表述为一个带同城约束的离散分配模型："
        "min sum_j [ C_rdc(j) * x_j ] + sum_i sum_j [ C_del(i,j) * z_ij ]。"
    )
    add_body_paragraph(
        doc,
        "约束包括：sum_j x_j = P；sum_j z_ij = 1（每个门店恰好分配一个RDC）；z_ij <= x_j；"
        "若 city(i) != city(j)，则 z_ij = 0。"
    )
    add_body_paragraph(
        doc,
        "程序对每个 P 值枚举合格 RDC 组合，并把门店分配给同城内配送折现成本最低的开放 RDC，从而得到该 P 下的最优网络方案。"
    )

    doc.add_heading("3.5 联合求解策略与最终评价指标", level=2)
    add_body_paragraph(
        doc,
        "为将新增门店收益与配送网络成本联动起来，程序采用“第二阶段组合 - 第三阶段网络子问题”嵌套求解。"
        "具体地，对每个新增门店组合 S，先由第二阶段给出其修正后收益，再调用第三阶段子问题求该组合对应的最优网络方案。"
    )
    add_body_paragraph(
        doc,
        "当前实现中，用于跨组合比较的联合指标定义为：Obj(S) = sum_i in S AdjNPV_i - DeliveryCost(S)。"
        "也就是说，RDC 生命周期成本在每个 P 子问题内被最小化并作为结果输出，但跨组合排序时使用的是“新增门店修正收益 - 配送成本”指标。"
    )
    add_body_paragraph(
        doc,
        "在报告展示层面，仍同步输出 TotalCost = RDCcost + DeliveryCost、平均配送距离、ROI、开放RDC清单与门店分配结果，"
        "以保证决策解释完整。"
    )
    add_body_paragraph(
        doc,
        "此外，程序在联合搜索中使用基于第二阶段收益上界的剪枝：若某组合的 Stage-2 调整后总 NPV 已不可能超过当前最好联合指标，"
        "则直接跳过其第三阶段求解，从而降低计算量。"
    )

    doc.add_heading("3.6 求解流程小结", level=2)
    add_body_paragraph(
        doc,
        "综上，整个模型的求解流程可概括为：先对新增门店做财务筛选，再对通过门店做带蚕食效应的组合优化，"
        "随后在每个可行组合上求解RDC选址与分配子问题，最后按联合指标选出推荐方案。"
    )
    add_body_paragraph(
        doc,
        "该框架兼顾了财务可行性、空间竞争与物流网络成本三类因素，适合本项目中“新增门店 + 既有门店 + RDC联动”的实际决策场景。"
    )

    doc.add_heading("四、结果的解读和可视化", level=1)
    add_body_paragraph(doc, "本节展示清理后最终三表口径下的双城结果，并给出选址与财务运营解释。")

    add_table(
        doc,
        counter.table("双城优化结果总览"),
        ["城市", "新增候选", "Stage1通过", "Stage2入选", "最优P", "RDC成本", "配送成本", "总成本", "ROI"],
        _result_summary_rows(res_bj, res_hz),
    )

    add_figure(doc, figures["stage3_cost"], counter.figure("Stage-3 成本构成对比（RDC成本+配送成本）"))
    add_figure(doc, figures["best_p_objective"], counter.figure("双城最优P与联合目标值对比"))

    open_rows = _open_rdc_rows("beijing", res_bj) + _open_rdc_rows("hangzhou", res_hz)
    add_table(doc, counter.table("最优方案开放RDC清单"), ["城市", "RDC ID", "名称", "生命周期成本(万元PV)"], open_rows)

    add_body_paragraph(
        doc,
        "结果解读要点：北京与杭州在最优P与成本结构上存在差异，体现了网络规模、距离结构与RDC成本约束的联合作用。"
    )

    doc.add_heading("附录A：数据清理与验收明细", level=1)
    removed = cleanup_summary.get("removed_by_rdc", {}) or {}
    removed_rows = [[k, v] for k, v in sorted(removed.items())]
    if not removed_rows:
        removed_rows = [["无", 0]]
    add_table(doc, counter.table("旧RDC记录删除明细"), ["rdc_id", "删除行数"], removed_rows)

    missing_pairs = checks.get("existing_store_missing_pairs", []) or []
    if missing_pairs:
        rows = [[x.get("store_id"), ",".join(x.get("missing_rdcs", []))] for x in missing_pairs]
    else:
        rows = [["无", "无"]]
    add_table(doc, counter.table("既有门店同城距离缺失明细"), ["store_id", "缺失RDC"], rows)

    doc.add_heading("参考文献", level=1)
    refs = [
        "[1] 高德开放平台. Web服务API 路径规划（含行驶距离计算）[EB/OL]. https://lbs.amap.com/api/webservice/guide/api/direction, 2026-04-18.",
        "[2] 北京市人民政府. 本市货车通行政策是什么？[EB/OL]. https://www.beijing.gov.cn/hudong/bmwd/jsjbmyyt/jsjbjtgl/jsjbjtgldhc/202207/t20220704_2762945.html, 2026-04-18.",
        "[3] 杭州市人民政府. 关于加强国四及以下柴油货车通行管理的通告[EB/OL]. https://www.hangzhou.gov.cn/art/2024/12/3/art_1229063381_1847467.html, 2026-04-18.",
        "[4] Happy Monkey项目组. 优化模型与系统代码（optimizer.py, app.py）[Z]. 仓库内部文档, 2026.",
        "[5] Happy Monkey项目组. 最终三表数据（stores/rdcs/distances）[DB/OL]. 用户提供数据, 2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.style = doc.styles["Normal"]

    ensure_dir(out_doc.parent)
    doc.save(out_doc)


def write_metrics_json(payload: Dict[str, object], cleanup_summary: Dict[str, object], out_path: Path) -> None:
    res_bj = payload["result_beijing"]
    res_hz = payload["result_hangzhou"]

    data = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "cleanup_summary": cleanup_summary,
        "quality_checks": payload["quality_checks"],
        "beijing_summary": res_bj["summary"],
        "hangzhou_summary": res_hz["summary"],
        "beijing_best": res_bj["stage3"].get("best_scenario"),
        "hangzhou_best": res_hz["stage3"].get("best_scenario"),
    }
    ensure_dir(out_path.parent)
    out_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    default_base = Path("/Users/losika/Desktop/美团商业数据分析大赛")
    parser = argparse.ArgumentParser(description="Generate final Happy Monkey report with distance cleanup summary.")
    parser.add_argument("--stores", type=Path, default=default_base / "stores.csv")
    parser.add_argument("--rdcs", type=Path, default=default_base / "rdcs.csv")
    parser.add_argument("--distances", type=Path, default=default_base / "distances.csv")

    parser.add_argument("--clean-distances", action="store_true", help="Backup and clean distances.csv by valid rdc ids.")
    parser.add_argument("--cleanup-backup-reference", type=Path, default=None, help="Use an existing backup file to derive cleanup summary without re-cleaning.")

    parser.add_argument("--output-dir", type=Path, default=Path("output/doc"))
    parser.add_argument("--output-doc", type=Path, default=Path("output/doc/快乐猴联合选址优化报告_终稿.docx"))
    parser.add_argument("--output-metrics", type=Path, default=Path("output/doc/report_metrics_final.json"))
    parser.add_argument("--output-cleanup", type=Path, default=Path("output/doc/data_cleanup_summary.json"))
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    ensure_dir(args.output_dir)
    figures_dir = args.output_dir / "figures_final"

    cleanup_summary: Dict[str, object]
    if args.clean_distances:
        cleanup_summary = cleanup_distances(args.distances, args.rdcs)
    elif args.cleanup_backup_reference:
        cleanup_summary = derive_cleanup_summary_from_backup(args.cleanup_backup_reference, args.distances, args.rdcs)
    else:
        cleanup_summary = {
            "cleaned": False,
            "backup_path": None,
            "original_rows": None,
            "kept_rows": len(load_rows(args.distances)),
            "removed_rows": None,
            "removed_by_rdc": {},
            "cleaned_at": datetime.now().isoformat(timespec="seconds"),
        }

    payload = build_datasets(args.stores, args.rdcs, args.distances)
    figures = generate_figures(payload, figures_dir)

    build_docx(payload, figures, cleanup_summary, args.output_doc)
    write_metrics_json(payload, cleanup_summary, args.output_metrics)
    args.output_cleanup.write_text(json.dumps(cleanup_summary, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Report generated: {args.output_doc}")
    print(f"Figures dir: {figures_dir}")
    print(f"Metrics JSON: {args.output_metrics}")
    print(f"Cleanup summary JSON: {args.output_cleanup}")


if __name__ == "__main__":
    main()
